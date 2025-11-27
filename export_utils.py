from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Image as RLImage,
)
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER

from datetime import datetime
from io import BytesIO
import re
import os


class MeetingExporter:
    def __init__(self, header_image_path="college_header.jpg"):
        self.header_image_path = header_image_path

    # ------------------------------------------------------
    # Helper: Clean general text
    # ------------------------------------------------------
    @staticmethod
    def _sanitize(text: str) -> str:
        if not text:
            return ""
        text = re.sub(r"[\u2580-\u259F\u2500-\u257F\u25A0-\u25FF]+", " ", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    # ------------------------------------------------------
    # Helper: Clean action item task text
    # ------------------------------------------------------
    @staticmethod
    def _clean_action_text(text: str) -> str:
        if not text:
            return ""

        # Remove Speaker names
        text = re.sub(r"Speaker\s*\d*:?", "", text, flags=re.IGNORECASE)
        text = re.sub(r"Speaker\s*:", "", text, flags=re.IGNORECASE)

        # Remove filler noise
        text = re.sub(r"\b(ma|am|ok|done|almost done)\b", "", text, flags=re.IGNORECASE)

        # Remove stray unicode lines
        text = re.sub(r"[\u2500-\u259F]+", "", text)

        # Normalize spaces
        text = re.sub(r"\s+", " ", text)

        # Trim punctuation
        return text.strip(" .,-:")

    # ------------------------------------------------------
    # Build clean action-item sentence
    # ------------------------------------------------------
    def _format_action_sentence(self, task, responsible="", deadline=""):
        task = self._clean_action_text(task)
        responsible = self._sanitize(responsible)
        deadline = self._sanitize(deadline)

        if not task:
            return ""

        # remove "to " at beginning
        task = re.sub(r"^\s*to\s+", "", task, flags=re.IGNORECASE).rstrip(".")

        subject = responsible if responsible else "The concerned staff"

        sentence = f"{subject} will {task}"
        if deadline:
            sentence += f" by {deadline}"
        sentence = sentence.strip()

        if not sentence.endswith("."):
            sentence += "."

        return sentence[0].upper() + sentence[1:]

    # ------------------------------------------------------
    # Formal summary generator
    # ------------------------------------------------------
    def _formal_summary_from_text(self, summary):
        s = self._sanitize(summary)
        if not s:
            return ""

        parts = [p.strip() for p in re.split(r"(?<=[.!?])\s+", s) if p.strip()]
        clean = []

        for p in parts:
            if not re.search(r"[.!?]$", p):
                p += "."
            p = p[0].upper() + p[1:]
            clean.append(p)

        if len(clean) == 1:
            return f"The meeting was convened to discuss the following: {clean[0]}"

        return "The meeting was convened to discuss the following points. " + " ".join(clean)

    # ------------------------------------------------------
    # Header insertion (DOCX)
    # ------------------------------------------------------
    def _add_docx_header(self, doc):
        if self.header_image_path and os.path.exists(self.header_image_path):
            try:
                picture = doc.add_picture(self.header_image_path, width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass

    # ------------------------------------------------------
    # Header insertion (PDF)
    # ------------------------------------------------------
    def _add_pdf_header(self, story):
        if self.header_image_path and os.path.exists(self.header_image_path):
            try:
                img = RLImage(self.header_image_path, width=6.5 * inch)
                story.append(img)
                story.append(Spacer(1, 12))
            except:
                pass

    # ------------------------------------------------------
    # DOCX Export
    # ------------------------------------------------------
    def export_to_docx(self, meeting_data):
        doc = Document()

        # Header image
        self._add_docx_header(doc)

        # Title
        heading = doc.add_heading("Minutes of Meeting", level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        metadata = meeting_data.get("metadata", {})

        for label in ["title", "date", "time", "venue", "organizer", "recorder"]:
            if metadata.get(label):
                p = doc.add_paragraph()
                p.add_run(label.capitalize() + ": ").bold = True
                p.add_run(str(metadata[label]))

        doc.add_paragraph("----")

        # Attendees
        attendees = meeting_data.get("attendees", [])
        if attendees:
            doc.add_heading("Attendees", level=2)
            for a in attendees:
                name = self._sanitize(a.get("name", ""))
                role = self._sanitize(a.get("role", ""))
                text = f"{name} – {role}" if role else name
                doc.add_paragraph(text, style="List Bullet")
            doc.add_paragraph("----")

        # AGENDA — simple bullet points
        agenda = meeting_data.get("agenda", [])
        if agenda:
            doc.add_heading("Agenda", level=2)
            for item in agenda:
                title = self._sanitize(item.get("title", ""))
                if title:
                    doc.add_paragraph(f"• {title}", style="List Bullet")
            doc.add_paragraph("----")

        # Summary
        summary = meeting_data.get("summary", "")
        if summary:
            doc.add_heading("Discussion Summary", level=2)
            doc.add_paragraph(self._formal_summary_from_text(summary))
            doc.add_paragraph("----")

        # Decisions
        decisions = meeting_data.get("decisions", [])
        if decisions:
            doc.add_heading("Decisions", level=2)
            for d in decisions:
                doc.add_paragraph(self._sanitize(d), style="List Bullet")
            doc.add_paragraph("----")

        # Action Items
        actions = meeting_data.get("action_items", [])
        if actions:
            doc.add_heading("Action Items", level=2)
            for a in actions:
                sentence = self._format_action_sentence(
                    a.get("task", ""),
                    a.get("responsible", ""),
                    a.get("deadline", "")
                )
                if sentence:
                    doc.add_paragraph(f"• {sentence}")
            doc.add_paragraph("----")

        # Next Meeting
        next_m = meeting_data.get("next_meeting", {})
        if any(next_m.values()):
            doc.add_heading("Next Meeting", level=2)
            for key in ["date", "time", "venue", "agenda"]:
                if next_m.get(key):
                    p = doc.add_paragraph()
                    p.add_run(key.capitalize() + ": ").bold = True
                    p.add_run(str(next_m[key]))
            doc.add_paragraph("─" * 60)

        # Closing
        doc.add_heading("Closing Note", level=2)
        doc.add_paragraph(
            f"Meeting minutes generated on {datetime.now().strftime('%d/%m/%Y at %H:%M')}."
        )

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ------------------------------------------------------
    # PDF Export
    # ------------------------------------------------------
    def export_to_pdf(self, meeting_data):
        buffer = BytesIO()

        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=20,
        )

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="TitleStyle", fontSize=20, alignment=TA_CENTER))
        styles.add(ParagraphStyle(name="Heading", fontSize=14, spaceAfter=10))
        styles.add(ParagraphStyle(name="Body", fontSize=11, leading=14))

        story = []

        # header
        self._add_pdf_header(story)

        # Title
        story.append(Paragraph("Minutes of Meeting", styles["TitleStyle"]))
        story.append(Spacer(1, 20))

        # Metadata
        metadata = meeting_data.get("metadata", {})
        for label in ["title", "date", "time", "venue", "organizer", "recorder"]:
            if metadata.get(label):
                story.append(
                    Paragraph(
                        f"<b>{label.capitalize()}:</b> {self._sanitize(metadata[label])}",
                        styles["Body"],
                    )
                )
        story.append(Spacer(1, 12))
        story.append(Paragraph("----", styles["Body"]))
        story.append(Spacer(1, 12))

        # Attendees
        attendees = meeting_data.get("attendees", [])
        if attendees:
            story.append(Paragraph("Attendees", styles["Heading"]))
            for a in attendees:
                name = self._sanitize(a.get("name", ""))
                role = self._sanitize(a.get("role", ""))
                line = f"• {name} – {role}" if role else f"• {name}"
                story.append(Paragraph(line, styles["Body"]))
            story.append(Spacer(1, 12))
            story.append(Paragraph("----", styles["Body"]))
            story.append(Spacer(1, 12))

        # Agenda
        agenda = meeting_data.get("agenda", [])
        if agenda:
            story.append(Paragraph("Agenda", styles["Heading"]))
            for item in agenda:
                title = self._sanitize(item.get("title", ""))
                if title:
                    story.append(Paragraph(f"• {title}", styles["Body"]))
            story.append(Spacer(1, 12))
            story.append(Paragraph("----", styles["Body"]))
            story.append(Spacer(1, 12))

        # Summary
        summary = meeting_data.get("summary", "")
        if summary:
            story.append(Paragraph("Discussion Summary", styles["Heading"]))
            story.append(Paragraph(self._formal_summary_from_text(summary), styles["Body"]))
            story.append(Spacer(1, 12))
            story.append(Paragraph("----", styles["Body"]))
            story.append(Spacer(1, 12))

        # Decisions
        decisions = meeting_data.get("decisions", [])
        if decisions:
            story.append(Paragraph("Decisions", styles["Heading"]))
            for d in decisions:
                story.append(Paragraph(f"• {self._sanitize(d)}", styles["Body"]))
            story.append(Spacer(1, 12))
            story.append(Paragraph("----", styles["Body"]))
            story.append(Spacer(1, 12))

        # Action items
        items = meeting_data.get("action_items", [])
        if items:
            story.append(Paragraph("Action Items", styles["Heading"]))
            for a in items:
                sentence = self._format_action_sentence(
                    a.get("task", ""),
                    a.get("responsible", ""),
                    a.get("deadline", "")
                )
                if sentence:
                    story.append(Paragraph(f"• {sentence}", styles["Body"]))
            story.append(Spacer(1, 12))
            story.append(
                Paragraph("----", styles["Body"])
            )
            story.append(Spacer(1, 12))

        # Next meeting
        next_m = meeting_data.get("next_meeting", {})
        if any(next_m.values()):
            story.append(Paragraph("Next Meeting", styles["Heading"]))
            for key in ["date", "time", "venue", "agenda"]:
                if next_m.get(key):
                    story.append(
                        Paragraph(
                            f"<b>{key.capitalize()}:</b> {self._sanitize(next_m[key])}",
                            styles["Body"],
                        )
                    )
            story.append(Spacer(1, 12))
            story.append(Paragraph("─" * 80, styles["Body"]))
            story.append(Spacer(1, 12))

        # Closing
        story.append(Paragraph("Closing Note", styles["Heading"]))
        story.append(
            Paragraph(
                f"Meeting minutes generated on {datetime.now().strftime('%d/%m/%Y at %H:%M')}.",
                styles["Body"],
            )
        )

        doc.build(story)
        buffer.seek(0)
        return buffer
